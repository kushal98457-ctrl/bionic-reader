"""
DOCX extractor.

Uses python-docx to pull paragraph text (including text inside tables).
Formatting/styles are intentionally dropped — the bionic engine only
needs the raw words; layout is a job for a future "preserve formatting"
extension if ever needed.
"""

from typing import BinaryIO, Union


def extract_from_path(path: str) -> str:
    import docx  # python-docx

    document = docx.Document(path)
    return _extract_from_document(document)


def extract_from_bytes(data: bytes) -> str:
    import io
    import docx

    document = docx.Document(io.BytesIO(data))
    return _extract_from_document(document)


def _extract_from_document(document) -> str:
    parts = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def extract(source: Union[str, bytes, BinaryIO]) -> str:
    if isinstance(source, str):
        return extract_from_path(source)
    if isinstance(source, (bytes, bytearray)):
        return extract_from_bytes(bytes(source))
    return extract_from_bytes(source.read())
