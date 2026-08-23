import io

from app.extractors import docx_extractor, text_extractor


def test_text_extractor_passthrough():
    assert text_extractor.extract_text("hello world") == "hello world"


def test_docx_extractor_reads_paragraphs():
    import docx

    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    document.save(buf)

    extracted = docx_extractor.extract_from_bytes(buf.getvalue())
    assert "First paragraph." in extracted
    assert "Second paragraph." in extracted


def test_pdf_extractor_reads_text():
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from a generated PDF")
    data = doc.tobytes()
    doc.close()

    from app.extractors import pdf_extractor

    extracted = pdf_extractor.extract_from_bytes(data)
    assert "Hello from a generated PDF" in extracted
